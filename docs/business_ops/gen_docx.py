# -*- coding: utf-8 -*-
"""Generate .docx versions of the two business templates (dual-format delivery)."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ACCENT = RGBColor(0x25, 0x63, 0xEB)

def style_doc(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(__import__("docx").oxml.ns.qn("w:eastAsia"), "Microsoft YaHei")

def h(doc, text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p

def h2(doc, text):
    return doc.add_heading(text, level=2)

def para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p

def bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")

def table(doc, rows):
    t = doc.add_table(rows=0, cols=len(rows[0]))
    t.style = "Light Grid Accent 1"
    for row in rows:
        cells = t.add_row().cells
        for i, c in enumerate(row):
            cells[i].text = c
    return t

# ---------- Template 1: 对公收款说明模板 ----------
d1 = Document()
style_doc(d1)
title = d1.add_heading("对公收款说明模板（致 B2B 客户）", level=0)
para(d1, "收款方（乙方）：上海杜特企业管理咨询有限公司")
para(d1, "使用说明：填写所有〔〕占位项后，可直接复制以下文本发给客户。核心原则——所有经营收入走对公银行账户，不使用个人微信/支付宝收款，便于对账与开票。", italic=True)

h2(d1, "一、收款账户信息")
para(d1, "请贵司通过对公银行账户转账至以下账户（请勿使用个人网银、支付宝或微信个人账号付款，以便对账与开具发票）：")
table(d1, [
    ["项目", "内容"],
    ["户名", "上海杜特企业管理咨询有限公司"],
    ["统一社会信用代码（税号）", "〔填写 18 位统一社会信用代码〕"],
    ["开户银行", "上海农商银行陈行支行"],
    ["银行账号", "32434508010036375"],
    ["注册地址", "〔营业执照注册地址〕"],
    ["联系电话", "13636690529 / 13311602075 / 13901700712"],
])

h2(d1, "二、付款指引")
for t in [
    "请使用贵司对公账户转账，备注栏填写：〔合同编号 / 订单号 / 项目名称〕+ 款项性质（如“2026 年 X 月企业培训费”）。",
    "转账金额以双方确认的《合同》/《报价单》为准；如有疑问请先与下方联系人核对，勿自行拆分或多付。",
    "如贵司财务制度允许个人代付小额款项（报名费、资料费等），请个人付款时务必在附言注明公司全称 + 用途，否则无法对应开票。",
]:
    bullet(d1, t)

h2(d1, "三、发票开具")
for t in [
    "我司可开具：〔增值税普通发票 / 增值税专用发票〕，开票内容一般为“*咨询服务*服务费”或“*非学历教育培训*培训费”（以实际业务为准）。",
    "请付款后提供以下开票信息：单位全称、税号；开户行及账号、注册地址及电话（专票必需）；收票邮箱（电子发票）或收件地址（纸质）。",
    "我司在收到款项并核对无误后〔3–5 个工作日〕内开具并送达发票。增值税专用发票请务必提前确认信息准确，开错重开周期较长。",
]:
    bullet(d1, t)

h2(d1, "四、到账与确认")
for t in [
    "付款后请将银行电子回单（截图或 PDF）发给联系人，以便对账。",
    "我司确认到账后，即启动对应服务（培训排期 / 咨询交付 / 资料发送等）。",
    "因付款备注缺失导致无法对应挂账的，请主动联系，避免影响服务启动。",
]:
    bullet(d1, t)

h2(d1, "五、联系人")
bullet(d1, "杜先生：13636690529 / 13311602075")
bullet(d1, "范女士：13901700712")

# ---------- Template 2: 银行对公聚合码咨询清单 ----------
d2 = Document()
style_doc(d2)
d2.add_heading("银行对公聚合收款码 · 咨询清单", level=0)
para(d2, "目的：在“直接对公收款”思路下，给个人客户也提供扫码便利，同时资金直接进对公户、不经过微信/支付宝资金池。去开户行（或电话客户经理）逐条问清以下问题，记下每项答复，再对比 2–3 家银行后决定。", italic=True)

sections = {
    "一、基础资格": [
        "该产品是否支持小微企业 / 普通有限责任公司申请？是否需额外资质？",
        "资金是否直接清算到我司对公银行账户？（关键：不是进微信 / 支付宝余额再提现）",
        "申请是否需要重新开户，还是绑定现有对公户即可？",
    ],
    "二、资金与清算": [
        "清算模式：是“银行直清到对公户”还是“经第三方支付机构二次清算”？（优先选银行直清，资金权属最干净）",
        "到账周期：T+0 / T+1 / D+1？是否支持实时到账？",
        "是否支持退款及退款到账规则？",
    ],
    "三、费率与费用（重点砍价）": [
        "微信通道费率、支付宝通道费率分别是多少？是否有小微企业费率优惠 / 减免活动？",
        "是否有单笔 / 单日限额影响大额收款？",
        "除交易费率外，是否有：开通费、年费、设备费（扫码枪 / 播报音箱）、提现费、对账服务费？",
        "对公账户本身的管理费 / 网银年费是否能因该产品减免？（小微企业多可免）",
    ],
    "四、通道与体验": [
        "支持哪些扫码方式：微信、支付宝、银联云闪付、数字人民币？",
        "是否支持静态收款码（打印张贴）与动态收款码（按金额生成）？",
        "是否提供：收款语音播报、多店员 / 多门店管理、经营报表？",
    ],
    "五、技术对接（为 LDA 商店预留）": [
        "是否提供开放 API / 商户平台接口，可对接自有系统实现“支付成功 → 自动交付”？",
        "是否支持异步回调（支付通知）与订单号关联？",
        "若无 API，是否至少支持流水导出（Excel / API）供我司系统对账？",
    ],
    "六、对账与限额": [
        "后台对账功能：按交易 / 按日 / 按门店汇总？能否导出？",
        "单笔、单日、单月收款限额是多少？可否申请提额？",
        "风控规则：是否有可疑交易拦截、冻结机制？申诉渠道？",
    ],
    "七、申请资料（提前备齐）": [
        "营业执照（三证合一）、法人身份证、开户许可证 / 基本户信息、经办人身份证",
        "经营场景证明（如门店照 / 线上店铺链接，视银行要求）",
        "预计申请时长：〔记下银行答复〕",
    ],
}
for title, items in sections.items():
    h2(d2, title)
    for t in items:
        bullet(d2, t)

import os
out1 = os.path.join(os.path.dirname(__file__), "对公收款说明模板.docx")
out2 = os.path.join(os.path.dirname(__file__), "银行对公聚合码咨询清单.docx")
d1.save(out1)
d2.save(out2)
print("saved", out1, out2)
