# -*- coding: utf-8 -*-
"""LDA 白皮书 Markdown -> Word (.docx) 转换器（零外部服务，纯 python-docx）。

用法：
  python scripts/whitepaper_md_to_docx.py <SRC.md> <OUT.docx>

支持元素：
  #/##/### 标题、| 表格 |、**加粗** 内联、> 引用块、-/* 无序与 1. 有序列表、--- 分隔。
白皮书中所有对比矩阵均为标准 markdown 表格，本转换器按表头/分隔行识别并渲染为带边框表格。
"""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def parse_inline(paragraph, text):
    """把 **加粗** 渲染为 bold run，其余原样。"""
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def split_blocks(lines):
    """把 markdown 行流切成块序列：('h', level, text) / ('table', [rows]) /
    ('quote', text) / ('list', [raw_lines]) / ('p', text)。"""
    blocks = []
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        if line.strip() == "":
            i += 1
            continue
        # 标题
        m = re.match(r"^(#{1,3})\s+(.*)$", line)
        if m:
            blocks.append(("h", len(m.group(1)), m.group(2).strip()))
            i += 1
            continue
        # 表格：当前行以 | 起，且下一行是分隔行 |---|
        if (line.lstrip().startswith("|") and i + 1 < n
                and re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[i + 1])):
            tbl = [[c.strip() for c in line.strip().strip("|").split("|")]]
            i += 2
            while i < n and lines[i].lstrip().startswith("|"):
                tbl.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            blocks.append(("table", tbl))
            continue
        # 引用块
        if line.lstrip().startswith(">"):
            buf = []
            while i < n and lines[i].lstrip().startswith(">"):
                buf.append(lines[i].lstrip()[1:].strip())
                i += 1
            blocks.append(("quote", " ".join(buf)))
            continue
        # 列表（无序 -/* 或有序 1.）
        if re.match(r"^\s*([-*]|\d+\.)\s+", line):
            buf = []
            while i < n and re.match(r"^\s*([-*]|\d+\.)\s+", lines[i]):
                buf.append(lines[i])
                i += 1
            blocks.append(("list", buf))
            continue
        # 分隔符
        if re.match(r"^---+\s*$", line):
            i += 1
            continue
        # 普通段落：合并后续连续正文行
        buf = [line]
        i += 1
        while (i < n and lines[i].strip() != ""
               and not lines[i].lstrip().startswith(("#", "|", ">"))
               and not re.match(r"^\s*([-*]|\d+\.)\s+", lines[i])
               and not re.match(r"^---+\s*$", lines[i])):
            buf.append(lines[i])
            i += 1
        blocks.append(("p", " ".join(x.strip() for x in buf)))
    return blocks


def render(doc, blocks):
    for blk in blocks:
        kind = blk[0]
        if kind == "h":
            _, lvl, text = blk
            style = {1: "Title", 2: "Heading 1", 3: "Heading 2"}.get(lvl, "Normal")
            p = doc.add_paragraph(style=style)
            parse_inline(p, text)
        elif kind == "p":
            p = doc.add_paragraph()
            parse_inline(p, blk[1])
        elif kind == "quote":
            p = doc.add_paragraph(style="Intense Quote")
            parse_inline(p, blk[1])
        elif kind == "list":
            ordered = bool(re.match(r"^\s*\d+\.\s+", blk[1][0]))
            for raw in blk[1]:
                m = re.match(r"^\s*(?:[-\*]|\d+\.)\s+(.*)$", raw)
                content = m.group(1) if m else raw
                p = doc.add_paragraph(style="List Number" if ordered else "List Bullet")
                parse_inline(p, content)
        elif kind == "table":
            rows = blk[1]
            t = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
            t.style = "Table Grid"
            for r_i, row in enumerate(rows):
                cells = t.rows[r_i].cells
                for c_i, val in enumerate(row):
                    if c_i >= len(cells):
                        continue
                    para = cells[c_i].paragraphs[0]
                    parse_inline(para, val)
                    if r_i == 0:  # 表头加粗
                        for run in para.runs:
                            run.bold = True
    return doc


def main():
    src = sys.argv[1] if len(sys.argv) > 1 else \
        "D:/agent_LDA/docs/lda_technical_whitepaper_2026-09-05.md"
    out = sys.argv[2] if len(sys.argv) > 2 else \
        "D:/agent_LDA/lda/lda_webui/static/lda_whitepaper_v0.9.40.docx"
    lines = open(src, encoding="utf-8").read().split("\n")
    doc = render(Document(), split_blocks(lines))
    doc.save(out)
    print("WROTE", out)


if __name__ == "__main__":
    main()
